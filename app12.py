import streamlit as st
import pandas as pd
import numpy as np
from langchain_huggingface import HuggingFaceEndpoint
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from docx import Document
import io
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
HF_API_KEY = os.getenv("HF_API_KEY")

# Verify API key
if not HF_API_KEY:
    st.error("HF_API_KEY not found in .env file. Please set it.")
    st.stop()

# Initialize Hugging Face model
llm = HuggingFaceEndpoint(
    repo_id="mistralai/Mixtral-8x7B-Instruct-v0.1",
    huggingfacehub_api_token=HF_API_KEY,
    temperature=0.7,
    max_new_tokens=500
)

# Define the prompt template
prompt_template = """
You are a data quality expert. Given the following data quality issues for a CSV file:
{issues}
Provide a detailed explanation of each issue, including:
1. What the issue is and why it matters.
2. Suggested fixes (e.g., drop rows, impute values, remove duplicates).
3. Potential impact of applying the fix.
Return the response in clear, concise language suitable for a data scientist.
"""
prompt = PromptTemplate(template=prompt_template, input_variables=["issues"])
llm_chain = LLMChain(llm=llm, prompt=prompt)

# Function to inspect data quality
def inspect_data(df):
    issues = {}
    
    # Check for missing values
    missing = df.isnull().sum()
    missing = missing[missing > 0]
    if not missing.empty:
        issues["Missing Values"] = missing.to_dict()
    
    # Check for duplicates
    duplicates = df.duplicated().sum()
    if duplicates > 0:
        issues["Duplicate Rows"] = duplicates
    
    # Check for outliers (IQR method for numeric columns)
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    outliers = {}
    for col in numeric_cols:
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        outlier_count = ((df[col] < (Q1 - 1.5 * IQR)) | (df[col] > (Q3 + 1.5 * IQR))).sum()
        if outlier_count > 0:
            outliers[col] = outlier_count
    if outliers:
        issues["Outliers"] = outliers
    
    # Check for schema mismatch
    dtypes = df.dtypes.to_dict()
    schema_issues = {}
    for col, dtype in dtypes.items():
        if df[col].apply(lambda x: isinstance(x, str) if pd.notnull(x) else False).all() and dtype != "object":
            schema_issues[col] = f"Expected string, found {dtype}"
        elif df[col].apply(lambda x: isinstance(x, (int, float)) if pd.notnull(x) else False).all() and dtype not in ["int64", "float64"]:
            schema_issues[col] = f"Expected numeric, found {dtype}"
    if schema_issues:
        issues["Schema Mismatch"] = schema_issues
    
    return issues

# Function to clean data
def clean_data(df, fix_missing, fix_duplicates, fix_outliers):
    cleaned_df = df.copy()
    logs = []
    
    # Handle missing values
    if fix_missing == "Drop":
        initial_rows = len(cleaned_df)
        cleaned_df = cleaned_df.dropna()
        logs.append(f"Dropped {initial_rows - len(cleaned_df)} rows with missing values.")
    elif fix_missing == "Impute":
        for col in cleaned_df.columns:
            if cleaned_df[col].isnull().any():
                if cleaned_df[col].dtype in ["int64", "float64"]:
                    mean_val = cleaned_df[col].mean()
                    cleaned_df[col] = cleaned_df[col].fillna(mean_val)
                    logs.append(f"Imputed missing values in {col} with mean ({mean_val:.2f}).")
                else:
                    mode_val = cleaned_df[col].mode()[0]
                    cleaned_df[col] = cleaned_df[col].fillna(mode_val)
                    logs.append(f"Imputed missing values in {col} with mode ({mode_val}).")
    
    # Handle duplicates
    if fix_duplicates:
        initial_rows = len(cleaned_df)
        cleaned_df = cleaned_df.drop_duplicates()
        logs.append(f"Removed {initial_rows - len(cleaned_df)} duplicate rows.")
    
    # Handle outliers
    if fix_outliers:
        for col in cleaned_df.select_dtypes(include=[np.number]).columns:
            Q1 = cleaned_df[col].quantile(0.25)
            Q3 = cleaned_df[col].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            initial_rows = len(cleaned_df)
            cleaned_df = cleaned_df[(cleaned_df[col] >= lower_bound) & (cleaned_df[col] <= upper_bound)]
            logs.append(f"Removed {initial_rows - len(cleaned_df)} rows with outliers in {col}.")
    
    return cleaned_df, logs

# Function to generate Word report
def generate_report(issues, llm_response, logs):
    doc = Document()
    doc.add_heading("Data Quality Report", 0)
    
    doc.add_heading("Detected Issues", level=1)
    for issue_type, details in issues.items():
        doc.add_heading(issue_type, level=2)
        doc.add_paragraph(str(details))
    
    doc.add_heading("AI Analysis", level=1)
    doc.add_paragraph(llm_response)
    
    doc.add_heading("Cleaning Logs", level=1)
    for log in logs:
        doc.add_paragraph(log)
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# Streamlit UI
st.title("Agentic AI Data Quality Checker (Hugging Face)")
st.write("Upload a CSV file to inspect and clean data quality issues using a free Hugging Face model.")

# File uploader
uploaded_file = st.file_uploader("Choose a CSV file", type="csv")

if uploaded_file is not None:
    # Read CSV
    df = pd.read_csv(uploaded_file)
    st.write("### Data Preview")
    st.dataframe(df.head())
    
    # Inspect data
    issues = inspect_data(df)
    
    if issues:
        st.write("### Detected Data Quality Issues")
        st.json(issues)
        
        # Get LLM explanations
        with st.spinner("Analyzing issues with Hugging Face model..."):
            llm_response = llm_chain.run(issues=str(issues))
            st.write("### AI Analysis and Suggestions")
            st.markdown(llm_response)
        
        # User options for fixes
        st.write("### Apply Fixes")
        fix_missing = st.selectbox("Handle missing values", ["None", "Drop", "Impute"])
        fix_duplicates = st.checkbox("Remove duplicates")
        fix_outliers = st.checkbox("Remove outliers")
        
        if st.button("Apply Fixes"):
            cleaned_df, logs = clean_data(df, fix_missing, fix_duplicates, fix_outliers)
            st.write("### Cleaning Logs")
            for log in logs:
                st.write(log)
            st.write("### Cleaned Data Preview")
            st.dataframe(cleaned_df.head())
            
            # Export cleaned data
            output_csv = io.StringIO()
            cleaned_df.to_csv(output_csv, index=False)
            st.download_button(
                label="Download Cleaned CSV",
                data=output_csv.getvalue(),
                file_name="cleaned_data.csv",
                mime="text/csv"
            )
            
            # Generate and download report
            report_data = generate_report(issues, llm_response, logs)
            st.download_button(
                label="Download Report (DOCX)",
                data=report_data,
                file_name="data_quality_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.write("No data quality issues detected!")